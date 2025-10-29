"""Microsoft Word document parser."""

from pathlib import Path
from .base import BaseParser, ParseResult


class DocxParser(BaseParser):
    """Microsoft Word 문서 파서 (python-docx 사용)"""

    def parse(self, file_path: Path) -> ParseResult:
        """
        DOCX 파일에서 텍스트를 추출합니다.

        Args:
            file_path: 파싱할 DOCX 파일 경로

        Returns:
            ParseResult: 파싱된 텍스트와 메타데이터
        """
        if not file_path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")

        try:
            try:
                from docx import Document
            except ImportError:
                raise ImportError(
                    "python-docx가 설치되지 않았습니다. 'pip install python-docx'로 설치하세요."
                )

            doc = Document(file_path)

            # 모든 단락에서 텍스트 추출
            paragraphs = [para.text for para in doc.paragraphs]
            content = '\n'.join(paragraphs)

            # 표(table)에서도 텍스트 추출 (옵션)
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    tables_text.append(row_text)

            if tables_text:
                content += '\n\n' + '\n'.join(tables_text)

            metadata = {
                'file_type': 'docx',
                'num_paragraphs': len(paragraphs),
                'num_tables': len(doc.tables),
                'file_size': file_path.stat().st_size,
            }

            return ParseResult(content=content, metadata=metadata)

        except Exception as e:
            raise ValueError(f"DOCX 파싱 실패: {file_path}") from e

    def supports(self, file_path: Path) -> bool:
        """DOCX 파일 확장자 지원 확인"""
        return file_path.suffix.lower() in self.supported_extensions

    @property
    def supported_extensions(self) -> list[str]:
        """지원하는 Word 문서 확장자"""
        return ['.docx']
